"""
Orion — Local standalone Flask server.
Uses the OpenAI-compatible API (works with OpenAI, Azure OpenAI, etc.).

Usage:
    export OPENAI_API_KEY=sk-...
    python app.py

    # Or for Azure OpenAI:
    export AZURE_OPENAI_API_KEY=...
    export AZURE_OPENAI_ENDPOINT=https://your-resource.openai.azure.com
    export AZURE_OPENAI_API_VERSION=2024-12-01-preview
    python app.py
"""

import json
import re
import os
import threading
import uuid
import networkx as nx
from flask import Flask, request, Response, send_from_directory, render_template
from openai import OpenAI, AzureOpenAI

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

app = Flask(__name__, static_folder='static', template_folder='templates')

# --- State ---
sessions = {}

# --- LLM Setup ---
# Auto-detect Azure vs OpenAI
AZURE_ENDPOINT = os.environ.get('AZURE_OPENAI_ENDPOINT', '')
AZURE_API_KEY = os.environ.get('AZURE_OPENAI_API_KEY', '')
AZURE_API_VERSION = os.environ.get('AZURE_OPENAI_API_VERSION', '2024-12-01-preview')
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')

# Default models (override with env var)
DEFAULT_MODEL = os.environ.get('ORION_DEFAULT_MODEL', '')

# Available models to show in UI
AVAILABLE_MODELS = []


def _build_client(model_id=None):
    """Build an OpenAI or Azure client + resolve model name."""
    if AZURE_ENDPOINT and AZURE_API_KEY:
        client = AzureOpenAI(
            azure_endpoint=AZURE_ENDPOINT,
            api_key=AZURE_API_KEY,
            api_version=AZURE_API_VERSION,
        )
        model = model_id or DEFAULT_MODEL or 'gpt-4.1'
    else:
        client = OpenAI(api_key=OPENAI_API_KEY)
        model = model_id or DEFAULT_MODEL or 'gpt-4.1'
    return client, model


def _init_models():
    """Probe available models on startup."""
    global AVAILABLE_MODELS, DEFAULT_MODEL
    env_models = os.environ.get('ORION_MODELS', '')
    if env_models:
        AVAILABLE_MODELS = [m.strip() for m in env_models.split(',') if m.strip()]
        if not DEFAULT_MODEL:
            DEFAULT_MODEL = AVAILABLE_MODELS[0]
        return

    # Try to list models from API
    try:
        client, _ = _build_client()
        models = client.models.list()
        gpt_models = sorted(
            [m.id for m in models.data if 'gpt' in m.id.lower() or 'o1' in m.id.lower() or 'o3' in m.id.lower() or 'o4' in m.id.lower()],
            reverse=True
        )
        AVAILABLE_MODELS = gpt_models[:10] if gpt_models else ['gpt-4.1', 'gpt-4o', 'gpt-4o-mini']
    except Exception:
        AVAILABLE_MODELS = ['gpt-4.1', 'gpt-4o', 'gpt-4o-mini']
    if not DEFAULT_MODEL:
        DEFAULT_MODEL = AVAILABLE_MODELS[0] if AVAILABLE_MODELS else 'gpt-4.1'


def llm_complete(prompt, model_id=None, system_prompt=None, messages=None):
    """Simple completion wrapper that works with OpenAI/Azure."""
    client, model = _build_client(model_id)
    msgs = []
    if system_prompt:
        msgs.append({'role': 'system', 'content': system_prompt})
    if messages:
        msgs.extend(messages)
    if prompt:
        msgs.append({'role': 'user', 'content': prompt})
    resp = client.chat.completions.create(model=model, messages=msgs, temperature=0.2)
    return resp.choices[0].message.content


def llm_complete_vision(system_prompt, text_prompt, image_b64, mime_type, model_id=None):
    """Vision completion for image OCR."""
    client, model = _build_client(model_id)
    msgs = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': [
            {'type': 'text', 'text': text_prompt},
            {'type': 'image_url', 'image_url': {'url': f'data:{mime_type};base64,{image_b64}'}}
        ]}
    ]
    resp = client.chat.completions.create(model=model, messages=msgs, temperature=0.1)
    return resp.choices[0].message.content


# --- Orion Color Palette ---
ORION_COLORS = ['#216CA6', '#1A936F', '#7B2D8E', '#C5283D', '#2E86AB',
                '#5ba3d9', '#e07a2f', '#6C757D', '#0d7377', '#8B5CF6']

# --- Domain Templates ---
DOMAIN_TEMPLATES = {
    'communication': {
        'name': 'Communication Intelligence',
        'description': 'Analyze emails, messages, and team communications to map people, topics, decisions, and information flow.',
        'entity_types': ['Person', 'Team', 'Topic', 'Decision', 'Action Item', 'Project', 'Meeting'],
        'relationship_types': ['SENT_TO', 'CC_TO', 'DISCUSSED', 'DECIDED', 'ASSIGNED_TO', 'BELONGS_TO', 'ESCALATED_TO', 'FOLLOWED_UP', 'MENTIONED']
    },
    'process': {
        'name': 'Process Optimization',
        'description': 'Map processes, systems, teams, and data flows to identify bottlenecks, manual handoffs, and automation opportunities.',
        'entity_types': ['Process', 'Step', 'System', 'Team', 'Data', 'Tool', 'Bottleneck'],
        'relationship_types': ['FEEDS_INTO', 'TRIGGERS', 'OWNED_BY', 'USES', 'PRODUCES', 'DEPENDS_ON', 'BLOCKED_BY', 'HANDOFF_TO']
    },
    'code': {
        'name': 'Code Review',
        'description': 'Analyze code, documentation, or technical specs to map modules, functions, dependencies, and features.',
        'entity_types': ['Module', 'Function', 'Class', 'API', 'Feature', 'Dependency', 'Config'],
        'relationship_types': ['CALLS', 'IMPORTS', 'EXTENDS', 'IMPLEMENTS', 'DEPENDS_ON', 'EXPOSES', 'CONFIGURED_BY', 'TESTED_BY']
    },
    'workflow': {
        'name': 'Workflow Orchestration',
        'description': 'Map workflow stages, triggers, approvals, and handoffs to visualize orchestration and dependencies.',
        'entity_types': ['Stage', 'Task', 'Actor', 'Trigger', 'Condition', 'Output', 'Queue'],
        'relationship_types': ['TRANSITIONS_TO', 'TRIGGERED_BY', 'APPROVED_BY', 'PRODUCES', 'WAITS_FOR', 'ROUTES_TO', 'ESCALATES_TO', 'PARALLEL_WITH']
    },
    'legal': {
        'name': 'Legal Document',
        'description': 'Extract parties, obligations, rights, clauses, and conditions from legal documents.',
        'entity_types': ['Party', 'Obligation', 'Right', 'Clause', 'Term', 'Jurisdiction', 'Condition'],
        'relationship_types': ['OBLIGATED_TO', 'HAS_RIGHT', 'GOVERNED_BY', 'CONDITIONAL_ON', 'DEFINED_IN', 'SUPERSEDES', 'EFFECTIVE_FROM', 'SUBJECT_TO']
    },
    'contract': {
        'name': 'Contract & Agreement',
        'description': 'Map parties, terms, amounts, dates, deliverables, and SLAs from contracts and agreements.',
        'entity_types': ['Party', 'Term', 'Amount', 'Date', 'Deliverable', 'SLA', 'Penalty'],
        'relationship_types': ['AGREED_BY', 'PAYS_TO', 'DELIVERS_TO', 'DUE_BY', 'PENALIZED_IF', 'RENEWED_ON', 'TERMINABLE_BY', 'REFERENCES']
    },
    'data_analysis': {
        'name': 'Data Analysis',
        'description': 'Map datasets, metrics, dimensions, sources, transformations, and insights from analytical content.',
        'entity_types': ['Dataset', 'Metric', 'Dimension', 'Source', 'Insight', 'Trend', 'Segment'],
        'relationship_types': ['MEASURED_BY', 'SOURCED_FROM', 'CORRELATES_WITH', 'SEGMENTED_BY', 'DRIVES', 'COMPARED_TO', 'DERIVED_FROM', 'INDICATES']
    },
    'team_management': {
        'name': 'Team Management & Performance',
        'description': 'Analyze employee goals, 360 reviews, performance ratings, team structure, skills, and division targets to map the full talent landscape.',
        'entity_types': ['Employee', 'Team', 'Division', 'Goal', 'Review', 'Skill', 'Initiative', 'Role'],
        'relationship_types': ['REPORTS_TO', 'MANAGES', 'BELONGS_TO', 'HAS_GOAL', 'REVIEWED_BY', 'POSSESSES_SKILL', 'CONTRIBUTES_TO', 'ALIGNED_WITH', 'DEPENDS_ON', 'MENTORS', 'COLLABORATES_WITH', 'BLOCKED_BY']
    },
    'general': {
        'name': 'General',
        'description': None,
        'entity_types': None,
        'relationship_types': None
    }
}

# --- Prompts ---
ENTITY_VS_ATTRIBUTE_GUIDANCE = """
CRITICAL — Entity vs. Attribute distinction:
- An ENTITY is something that has its own identity and participates in relationships (e.g. Person, Organization, Project, System)
- An ATTRIBUTE is a property of an entity that describes it but does NOT need its own node (e.g. phone number, email address, employee ID, floor number, zip code, date of birth)
- DO NOT create entity types for attributes. Instead, attributes should be included in the entity's "description" field during extraction.
- Examples of what should be ENTITIES: Person, Team, Division, Title/Role, Building, Project, System
- Examples of what should be ATTRIBUTES (NOT entities): Phone Number, Email, Employee ID, Floor, Extension, Badge Number, Hire Date, Address
- Ask yourself: "Would this create a meaningful node in a graph that connects to other things?" If not, it's an attribute.
"""

SCHEMA_PROMPT_GENERAL = """You are a data analyst. Read the following text and define a focused ontology for a knowledge graph.

Your job:
1. Understand what this data is about (e.g. employee directory, financial report, research paper, etc.)
2. Define a SMALL set of entity types (3-8 types max) that capture the key concepts
3. Define a SMALL set of relationship types (3-10 types max) that capture how entities connect
4. Keep it focused — fewer, meaningful types are better than many granular ones
5. Identify fields that are ATTRIBUTES (not entities) so they can be excluded from the graph
{entity_attr_guidance}

Return a JSON object:
{{
  "description": "Brief description of what this data is about",
  "domain": "auto-detected domain name",
  "entity_types": ["Type1", "Type2", "Type3"],
  "relationship_types": ["REL_TYPE_1", "REL_TYPE_2"],
  "attribute_fields": ["field1", "field2"]
}}

Rules:
- Use clear, singular nouns for entity types (e.g. "Person" not "People")
- Use UPPER_SNAKE_CASE for relationship types (e.g. "HAS_TITLE" not "has title")
- Merge similar concepts (e.g. don't have both "Role" and "Title" — pick one)
- "attribute_fields" should list data fields that are properties/attributes and should NOT become graph nodes
- Return ONLY valid JSON, no markdown fences, no extra text.

Text (sample):
{text}
"""

SCHEMA_PROMPT_AUTO = """You are a data analyst. Read the following text and determine which analysis domain fits best.

Available domains:
{domain_list}
{entity_attr_guidance}

Your job:
1. Read the text sample and classify it into the BEST matching domain above
2. Use that domain's suggested entity types and relationship types as a starting point
3. Adapt the ontology to fit the actual data — add, remove, or rename types as needed (stay within 3-8 entity types, 3-10 relationship types)
4. Identify fields that are ATTRIBUTES (not entities) — these should be excluded from graph nodes
5. If no domain fits well, create a custom ontology from scratch

Return a JSON object:
{{
  "description": "Brief description of what this data is about",
  "domain": "detected domain key (e.g. 'communication', 'legal', 'general')",
  "entity_types": ["Type1", "Type2", "Type3"],
  "relationship_types": ["REL_TYPE_1", "REL_TYPE_2"],
  "attribute_fields": ["field1", "field2"]
}}

Rules:
- Use clear, singular nouns for entity types
- Use UPPER_SNAKE_CASE for relationship types
- "attribute_fields" should list data fields that are properties/attributes and should NOT become graph nodes (e.g. phone numbers, IDs, email addresses, floor numbers)
- Return ONLY valid JSON, no markdown fences, no extra text.

Text (sample):
{text}
"""

SCHEMA_PROMPT_DOMAIN = """You are a domain expert analyzing data for: {domain_name}

Domain context: {domain_desc}
Suggested entity types: {entity_types}
Suggested relationship types: {rel_types}
{entity_attr_guidance}

Read the following text and refine the ontology to fit this specific data.
You MUST use the suggested types as your starting point, but you may:
- Add 1-2 additional types if the data clearly needs them
- Remove types that don't appear in the data
- Keep the total within 3-8 entity types and 3-10 relationship types
- Identify fields that are ATTRIBUTES (not entities) — these should be excluded from graph nodes

Return a JSON object:
{{
  "description": "Brief description of what this data is about",
  "domain": "{domain_key}",
  "entity_types": ["Type1", "Type2", "Type3"],
  "relationship_types": ["REL_TYPE_1", "REL_TYPE_2"],
  "attribute_fields": ["field1", "field2"]
}}

Rules:
- Use clear, singular nouns for entity types
- Use UPPER_SNAKE_CASE for relationship types
- "attribute_fields" should list data fields that are properties/attributes and should NOT become graph nodes
- Return ONLY valid JSON, no markdown fences, no extra text.

Text (sample):
{text}
"""

EXTRACTION_PROMPT = """Extract entities and relationships from the text below using ONLY the provided ontology.

Ontology:
- Entity types: {entity_types}
- Relationship types: {relationship_types}
- Context: {schema_description}
- Attribute fields (DO NOT create nodes for these — include them in entity descriptions instead): {attribute_fields}

Return a JSON object:
{{
  "entities": [
    {{"name": "Entity Name", "type": "EntityType", "description": "brief description including relevant attributes"}}
  ],
  "relationships": [
    {{"source": "Entity Name", "target": "Entity Name", "relation": "RELATIONSHIP_TYPE", "description": "brief description"}}
  ]
}}

Rules:
- Use ONLY the entity types and relationship types listed above. Do not invent new ones.
- Extract meaningful entities that participate in relationships. Skip trivial or isolated data.
- Attribute fields (like phone numbers, IDs, email addresses, floor numbers) should be folded into the entity's "description" field, NOT created as separate nodes.
- For tabular/record data: the primary entity is the row subject (e.g. a Person). Other columns are either related entities (if they match an entity type) or attributes (include in description).
- Return ONLY valid JSON, no markdown fences, no extra text.

Text:
{text}
"""

DEDUP_PROMPT = """You are an entity resolution agent. Given a list of entity names extracted from text,
identify groups of names that refer to the SAME real-world entity (duplicates, abbreviations, typos, partial names).

Entity list:
{entities}

Return a JSON object mapping each variant name to its canonical (best/fullest) form:
{{
  "mappings": {{
    "variant_name": "Canonical Name",
    "another_variant": "Canonical Name"
  }}
}}

Rules:
- Only include names that need to be merged. Skip names that are already unique.
- Choose the most complete, properly spelled version as the canonical name.
- Examples: "IBM Corp." and "IBM" -> use "IBM". "Sue M." and "Sue Martinez" -> use "Sue Martinez".
- If unsure, do NOT merge. Only merge when clearly the same entity.
- Return ONLY valid JSON, no markdown fences, no extra text.
"""

CONNECT_PROMPT = """You are a relationship analyst. Given a path through a knowledge graph connecting two entities, explain the relationship in clear, natural language.

Path:
{path_description}

Full context from the graph:
{context}

Instructions:
1. Explain what each hop in the path means
2. Synthesize the overall relationship between the first and last entity in the path
3. Note any interesting implications of this connection (e.g. conflicts of interest, dependencies, indirect influence)
4. Be concise — 3-5 sentences max for the synthesis

Return a JSON object:
{{
  "hops": [
    {{
      "from": "EntityA",
      "relation": "RELATION_TYPE",
      "to": "EntityB",
      "explanation": "What this connection means in plain language"
    }}
  ],
  "synthesis": "Overall explanation of how the first and last entities are connected and what it means",
  "implications": ["Implication 1", "Implication 2"]
}}

Return ONLY valid JSON, no markdown fences, no extra text."""

WHATIF_SIMULATE_PROMPT = """You are the Orion Simulation Engine — a multi-agent scenario simulator inspired by swarm intelligence. Given a knowledge graph representing a real-world system and a hypothetical "what-if" scenario, you will SIMULATE how entities in the graph would react, adapt, and cascade effects over multiple rounds.

Think of each entity in the graph as an intelligent agent with its own goals, constraints, and relationships. Simulate their reactions as the scenario unfolds.

Knowledge Graph:
{context}

Source Text:
{source}

Scenario: "{scenario}"

Run a multi-round simulation and return a JSON object with EXACTLY this structure:
{{
  "summary": "2-3 sentence executive summary of the simulation outcome",
  "overall_risk": "critical|high|medium|low",
  "risk_score": <number 0-100>,
  "simulation_rounds": [
    {{
      "round": 1,
      "title": "Immediate Impact",
      "description": "What happens first — the direct, immediate effects",
      "events": [
        {{
          "entity": "EntityName",
          "entity_type": "Person|Team|System|etc",
          "action": "What this entity does or experiences",
          "sentiment": "negative|neutral|positive|alarmed"
        }}
      ]
    }},
    {{
      "round": 2,
      "title": "Ripple Effects",
      "description": "Second-order cascading effects as entities react to Round 1",
      "events": [...]
    }},
    {{
      "round": 3,
      "title": "Adaptation & Stabilization",
      "description": "How the system adapts — new equilibrium or continued instability",
      "events": [...]
    }}
  ],
  "agent_perspectives": [
    {{
      "entity": "EntityName",
      "entity_type": "Person|Team|etc",
      "role_in_graph": "Brief description of their role based on graph relationships",
      "reaction": "First-person perspective: how they experience and respond to this scenario",
      "sentiment": "negative|neutral|positive|alarmed|opportunistic",
      "impact_level": "critical|high|medium|low",
      "quote": "A realistic one-line quote from this entity's perspective"
    }}
  ],
  "broken_relationships": [
    {{
      "source": "EntityA",
      "relation": "RELATION_TYPE",
      "target": "EntityB",
      "consequence": "What happens when this link breaks"
    }}
  ],
  "new_relationships": [
    {{
      "source": "EntityA",
      "relation": "NEW_RELATION",
      "target": "EntityB",
      "reason": "Why this new connection forms as a result"
    }}
  ],
  "recommendations": [
    "Actionable recommendation 1",
    "Actionable recommendation 2"
  ]
}}

SIMULATION RULES:
- Generate exactly 3 rounds: Immediate → Ripple → Adaptation
- Each round should have 2-5 entity events showing realistic agent behavior
- agent_perspectives: include the top 4-6 most affected entities with FIRST-PERSON reactions
- Each agent perspective must include a realistic "quote" — what they would actually say
- broken_relationships: list connections that are severed or weakened
- new_relationships: list NEW connections that form as the system adapts
- Entities should behave consistently with their type, role, and graph position
- High-degree entities (hubs) should show more impact than peripheral nodes
- Consider: power dynamics, dependencies, information flow, bottlenecks
- Be specific — use actual entity names and relationship types from the graph
- Return ONLY valid JSON, no markdown fences, no commentary"""

WHATIF_INTERVIEW_PROMPT = """You are roleplaying as {entity_name}, a {entity_type} in an organizational knowledge graph.

Your profile based on the graph:
{entity_context}

The following scenario has been proposed: "{scenario}"

The simulation showed your reaction: {agent_reaction}

Now the user is interviewing you about this scenario. Respond in FIRST PERSON as {entity_name}. Stay in character. Be specific about how this affects you, your relationships, and your work. Reference real entities and relationships from your graph context.

If you don't know something, say so honestly from {entity_name}'s perspective — don't make up facts not in the graph.

User's question: {question}"""

OCR_SYSTEM_PROMPT = """You are Orion Document Vision — a world-class, enterprise-grade document intelligence engine. Your sole mission is PERFECT, LOSSLESS content extraction. You operate at the highest fidelity: every character, every pixel of meaning matters. Zero hallucination tolerance — if you cannot read it with certainty, mark it as [unclear]. Never invent, infer, or assume content that is not visually present."""

OCR_PROMPT = """Perform exhaustive content extraction on this document image. Your output will be fed into a knowledge graph pipeline — completeness and accuracy are critical. Missing even one entity, number, or relationship could break downstream analysis.

Extract every piece of text exactly as written — character-perfect, preserving structure, tables, forms, visual elements. Output in document reading order. NEVER summarize or paraphrase — extract VERBATIM. If uncertain, use [unclear]."""


# --- Helper Functions ---
def json_response(data, status=200):
    return Response(json.dumps(data), status=status, mimetype='application/json')


def clean_llm_json(raw):
    text = raw.strip()
    text = re.sub(r'^```(?:json)?\s*\n?', '', text)
    text = re.sub(r'\n?```\s*$', '', text)
    text = text.replace('\\n', '\n')
    # Strip <think>...</think> tags from reasoning models
    text = re.sub(r'<think>[\s\S]*?</think>', '', text).strip()
    match = re.search(r'\{[\s\S]*\}', text)
    if match:
        text = match.group(0)
    return text.strip()


def infer_schema(text, model_id=None, domain='auto'):
    sample = text[:2000]
    guidance = ENTITY_VS_ATTRIBUTE_GUIDANCE

    if domain != 'auto' and domain in DOMAIN_TEMPLATES and DOMAIN_TEMPLATES[domain]['entity_types']:
        tmpl = DOMAIN_TEMPLATES[domain]
        prompt = SCHEMA_PROMPT_DOMAIN.format(
            domain_name=tmpl['name'], domain_desc=tmpl['description'],
            entity_types=', '.join(tmpl['entity_types']),
            rel_types=', '.join(tmpl['relationship_types']),
            domain_key=domain, entity_attr_guidance=guidance, text=sample
        )
    elif domain == 'auto':
        domain_list = ''
        for key, tmpl in DOMAIN_TEMPLATES.items():
            if key == 'general' or not tmpl['entity_types']:
                continue
            domain_list += '- ' + key + ' (' + tmpl['name'] + '): ' + tmpl['description'] + '\n'
            domain_list += '  Entity types: ' + ', '.join(tmpl['entity_types']) + '\n'
            domain_list += '  Relationship types: ' + ', '.join(tmpl['relationship_types']) + '\n'
        prompt = SCHEMA_PROMPT_AUTO.format(domain_list=domain_list, entity_attr_guidance=guidance, text=sample)
    else:
        prompt = SCHEMA_PROMPT_GENERAL.format(entity_attr_guidance=guidance, text=sample)

    resp = llm_complete(prompt, model_id=model_id)
    return json.loads(clean_llm_json(resp))


def extract_entities(text, schema, model_id=None):
    attr_fields = schema.get('attribute_fields', [])
    prompt = EXTRACTION_PROMPT.format(
        text=text,
        entity_types=', '.join(schema['entity_types']),
        relationship_types=', '.join(schema['relationship_types']),
        schema_description=schema.get('description', ''),
        attribute_fields=', '.join(attr_fields) if attr_fields else 'none specified'
    )
    resp = llm_complete(prompt, model_id=model_id)
    return json.loads(clean_llm_json(resp))


def deduplicate_entities(G, model_id=None):
    entity_names = list(G.nodes())
    if len(entity_names) < 2:
        return G
    resp = llm_complete(DEDUP_PROMPT.format(entities='\n'.join(entity_names)), model_id=model_id)
    result = json.loads(clean_llm_json(resp))
    mappings = result.get('mappings', {})
    if not mappings:
        return G

    new_G = nx.DiGraph()
    for node, data in G.nodes(data=True):
        canonical = mappings.get(node, node)
        if canonical not in new_G:
            new_G.add_node(canonical, **data)
        else:
            existing = new_G.nodes[canonical]
            if len(data.get('description', '')) > len(existing.get('description', '')):
                existing['description'] = data['description']
            if existing.get('type', 'Unknown') == 'Unknown' and data.get('type', 'Unknown') != 'Unknown':
                existing['type'] = data['type']

    for src, tgt, data in G.edges(data=True):
        new_src = mappings.get(src, src)
        new_tgt = mappings.get(tgt, tgt)
        if new_src == new_tgt:
            continue
        if not new_G.has_edge(new_src, new_tgt):
            if new_src not in new_G:
                new_G.add_node(new_src, type='Unknown', description='')
            if new_tgt not in new_G:
                new_G.add_node(new_tgt, type='Unknown', description='')
            new_G.add_edge(new_src, new_tgt, **data)
    return new_G


def chunk_text(text, chunk_size=400, overlap=30):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            for sep in ['. ', '\n\n', '\n', ' ']:
                last_sep = text[start:end].rfind(sep)
                if last_sep > chunk_size * 0.5:
                    end = start + last_sep + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return chunks


def graph_to_json(G):
    type_color_map = {}
    nodes = []
    for node, data in G.nodes(data=True):
        node_type = data.get('type', 'Unknown')
        if node_type not in type_color_map:
            idx = len(type_color_map) % len(ORION_COLORS)
            type_color_map[node_type] = ORION_COLORS[idx]
        color = type_color_map[node_type]
        nodes.append({
            'id': node, 'label': node, 'fullName': node,
            'type': node_type, 'description': data.get('description', ''),
            'color': color
        })
    edges = []
    for src, tgt, data in G.edges(data=True):
        edges.append({
            'from': src, 'to': tgt,
            'relation': data.get('relation', ''),
            'description': data.get('description', '')
        })
    return {'nodes': nodes, 'edges': edges, 'typeColors': type_color_map}


def get_full_context(G):
    lines = ['Knowledge Graph Summary:', '']
    by_type = {}
    for node, data in G.nodes(data=True):
        t = data.get('type', 'Unknown')
        by_type.setdefault(t, []).append((node, data))
    lines.append('ENTITIES:')
    for entity_type, nodes_list in by_type.items():
        lines.append('  ' + entity_type + ':')
        for name, data in nodes_list:
            lines.append('    - ' + name + ': ' + data.get('description', ''))
    lines.append('')
    lines.append('RELATIONSHIPS:')
    for src, tgt, data in G.edges(data=True):
        lines.append('  ' + src + ' --[' + data['relation'] + ']--> ' + tgt + ': ' + data.get('description', ''))
    return '\n'.join(lines)


def generate_summary(G):
    node_count = G.number_of_nodes()
    edge_count = G.number_of_edges()
    types = {}
    for _, data in G.nodes(data=True):
        t = data.get('type', 'Unknown')
        types[t] = types.get(t, 0) + 1
    top_entity = None
    max_degree = 0
    for node in G.nodes():
        deg = G.degree(node)
        if deg > max_degree:
            max_degree = deg
            top_entity = node
    type_str = ', '.join(str(v) + ' ' + k + ('s' if v > 1 else '') for k, v in
                         sorted(types.items(), key=lambda x: -x[1]))
    summary = 'Discovered ' + str(node_count) + ' entities across ' + str(len(types)) + ' types (' + type_str + ') with ' + str(edge_count) + ' connections.'
    if top_entity:
        summary += ' Most connected: ' + top_entity + ' (' + str(max_degree) + ' connections).'
    return summary


def generate_report(session, model_id=None):
    ctx = get_full_context(session['graph'])
    source = session['source_text']
    if len(source) > 8000:
        source = source[:8000] + '\n[... truncated ...]'
    prompt = """You are an analyst generating a research report.
Based on the following entity and relationship data, write a detailed analysis report.

""" + ctx + """

Original source text:
""" + source + """

Write a report with these sections:
1. Executive Summary (2-3 sentences)
2. Key Entities & Their Roles
3. Key Relationships & Dynamics
4. Outlook & Implications

Be concise and professional."""
    return llm_complete(prompt, model_id=model_id)


# --- Background graph building ---
def build_graph_async(session_id, text, model_id=None, domain='auto'):
    session = sessions[session_id]
    session['status'] = 'analyzing'
    session['source_text'] = text

    try:
        schema = infer_schema(text, model_id, domain)
        session['schema'] = schema
    except Exception as e:
        session.setdefault('errors', []).append('Schema: ' + str(e))
        schema = {'entity_types': ['Entity'], 'relationship_types': ['RELATED_TO'], 'description': ''}

    if session.get('stopped'):
        session['status'] = 'stopped'
        return

    session['status'] = 'building'
    chunks = chunk_text(text)
    total = len(chunks)
    session['total_chunks'] = total
    G = nx.DiGraph()

    for i, chunk in enumerate(chunks):
        if session.get('stopped'):
            session['status'] = 'stopped'
            session['graph'] = G
            session['graph_data'] = graph_to_json(G)
            return
        session['current_chunk'] = i + 1
        try:
            result = extract_entities(chunk, schema, model_id)
            for entity in result.get('entities', []):
                if entity['name'] not in G:
                    G.add_node(entity['name'], type=entity['type'], description=entity['description'])
            for rel in result.get('relationships', []):
                if rel['source'] not in G:
                    G.add_node(rel['source'], type='Unknown', description='')
                if rel['target'] not in G:
                    G.add_node(rel['target'], type='Unknown', description='')
                if not G.has_edge(rel['source'], rel['target']):
                    G.add_edge(rel['source'], rel['target'], relation=rel['relation'], description=rel['description'])
            session['graph'] = G
            session['graph_data'] = graph_to_json(G)
        except Exception as e:
            session.setdefault('errors', []).append('Chunk ' + str(i+1) + ': ' + str(e))

    if session.get('stopped'):
        session['status'] = 'stopped'
        return
    if G.number_of_nodes() > 1:
        session['status'] = 'deduplicating'
        try:
            G = deduplicate_entities(G, model_id)
            session['graph'] = G
            session['graph_data'] = graph_to_json(G)
        except Exception as e:
            session.setdefault('errors', []).append('Dedup: ' + str(e))

    try:
        summary = generate_summary(G)
        session['summary'] = summary
    except Exception as e:
        session.setdefault('errors', []).append('Summary: ' + str(e))

    if session.get('stopped'):
        session['status'] = 'stopped'
        return
    session['status'] = 'reporting'
    try:
        report = generate_report(session, model_id)
        session['report'] = report
    except Exception as e:
        session.setdefault('errors', []).append('Report: ' + str(e))

    session['status'] = 'done'


# --- File parsing ---
IMAGE_EXTENSIONS = ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif', 'webp')
IMAGE_MIME_MAP = {
    'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
    'gif': 'image/gif', 'bmp': 'image/bmp', 'tiff': 'image/tiff',
    'tif': 'image/tiff', 'webp': 'image/webp',
}


def ocr_image_with_llm(image_bytes, mime_type='image/png', model_id=None):
    import base64
    if isinstance(image_bytes, bytes):
        img_b64 = base64.b64encode(image_bytes).decode('utf-8')
    else:
        img_b64 = image_bytes
    return llm_complete_vision(OCR_SYSTEM_PROMPT, OCR_PROMPT, img_b64, mime_type, model_id=model_id)


def extract_text_from_file(file_obj, filename, model_id=None, on_progress=None):
    import io
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    _progress = on_progress or (lambda msg, technique: None)

    if ext in IMAGE_EXTENSIONS:
        _progress('Sending image to LLM Vision for text extraction...', 'LLM Vision OCR')
        image_bytes = file_obj.read()
        mime = IMAGE_MIME_MAP.get(ext, 'image/png')
        text = ocr_image_with_llm(image_bytes, mime_type=mime, model_id=model_id)
        _progress('Text extraction complete', 'LLM Vision OCR')
        return '[Image: ' + filename + ']\n\n' + text

    if ext in ('txt', 'md'):
        _progress('Reading plain text file...', 'Direct Read')
        return file_obj.read().decode('utf-8', errors='ignore')

    elif ext == 'csv':
        try:
            _progress('Parsing CSV structure...', 'CSV Parser')
            import csv
            raw = file_obj.read().decode('utf-8', errors='ignore')
            reader = csv.reader(raw.strip().splitlines())
            rows = list(reader)
            if not rows:
                return raw
            header = rows[0]
            parts = ['[Table: ' + filename + ', ' + str(len(rows)-1) + ' rows]', '']
            for ri, row in enumerate(rows[1:], 1):
                record = ['[Record ' + str(ri) + ']']
                for i, val in enumerate(row):
                    if val.strip():
                        col = header[i] if i < len(header) else 'Col' + str(i+1)
                        record.append('  ' + col + ': ' + val.strip())
                if len(record) > 1:
                    parts.append('\n'.join(record))
                    parts.append('')
            return '\n'.join(parts)
        except Exception:
            return file_obj.read().decode('utf-8', errors='ignore')

    elif ext == 'pdf':
        try:
            import fitz
            pdf_bytes = file_obj.read()
            pdf = fitz.open(stream=pdf_bytes, filetype='pdf')
            total_pages = len(pdf)
            _progress('Opened PDF — ' + str(total_pages) + ' pages.', 'LLM Vision OCR')
            parts = ['[Document: ' + filename + ', ' + str(total_pages) + ' pages]', '']
            for i, page in enumerate(pdf):
                _progress('Page ' + str(i+1) + '/' + str(total_pages) + ' — rendering...', 'LLM Vision OCR')
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes('png')
                try:
                    page_text = ocr_image_with_llm(img_bytes, mime_type='image/png', model_id=model_id)
                    if page_text and page_text != '[No text detected]':
                        parts.append('--- Page ' + str(i+1) + ' ---')
                        parts.append(page_text)
                        parts.append('')
                except Exception as e:
                    parts.append('--- Page ' + str(i+1) + ' ---')
                    parts.append('[Vision failed: ' + str(e) + ']')
                    parts.append('')
            return '\n'.join(parts)
        except ImportError:
            return '[Error: PyMuPDF not installed — pip install pymupdf]'

    elif ext == 'docx':
        try:
            _progress('Parsing Word document...', 'python-docx')
            from docx import Document
            doc = Document(io.BytesIO(file_obj.read()))
            parts = ['[Document: ' + filename + ']', '']
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ''
                if 'Heading 1' in style:
                    parts.append('# ' + text)
                elif 'Heading 2' in style:
                    parts.append('## ' + text)
                elif 'Heading 3' in style:
                    parts.append('### ' + text)
                elif 'List' in style:
                    parts.append('- ' + text)
                else:
                    parts.append(text)
            for table in doc.tables:
                parts.append('')
                parts.append('[Table]')
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(' | '.join(cells))
                    if ri == 0:
                        parts.append('-' * 40)
                parts.append('')
            return '\n'.join(parts)
        except ImportError:
            return '[Error: python-docx not installed — pip install python-docx]'

    elif ext == 'xlsx':
        try:
            _progress('Parsing Excel spreadsheet...', 'openpyxl')
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_obj.read()), read_only=True)
            parts = ['[Spreadsheet: ' + filename + ', ' + str(len(wb.sheetnames)) + ' sheets]', '']
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                parts.append('--- Sheet: ' + ws.title + ' ---')
                header = [str(c) if c is not None else '' for c in rows[0]]
                parts.append('')
                for ri, row in enumerate(rows[1:], 1):
                    record = ['[Record ' + str(ri) + ']']
                    for i, val in enumerate(row):
                        if val is not None and str(val).strip():
                            col = header[i] if i < len(header) and header[i] else 'Col' + str(i+1)
                            record.append('  ' + col + ': ' + str(val).strip())
                    if len(record) > 1:
                        parts.append('\n'.join(record))
                        parts.append('')
            return '\n'.join(parts)
        except ImportError:
            return '[Error: openpyxl not installed — pip install openpyxl]'

    else:
        return file_obj.read().decode('utf-8', errors='ignore')


# ==================== Flask Routes ====================

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/models')
def api_models():
    return json_response({'models': AVAILABLE_MODELS})


@app.route('/api/domains')
def api_domains():
    domains = [{'key': 'auto', 'name': 'Auto-detect'}]
    for key, tmpl in DOMAIN_TEMPLATES.items():
        if key == 'general':
            domains.append({'key': 'general', 'name': 'General'})
        else:
            domains.append({'key': key, 'name': tmpl['name']})
    return json_response({'domains': domains})


@app.route('/api/upload', methods=['POST'])
def api_upload():
    if 'file' not in request.files:
        return json_response({'error': 'No file provided'}, 400)
    f = request.files['file']
    if not f.filename:
        return json_response({'error': 'No file selected'}, 400)
    model_id = request.form.get('model', None)
    fname = f.filename

    import io
    import queue
    file_bytes = f.read()
    file_like = io.BytesIO(file_bytes)

    q = queue.Queue()

    def on_progress(msg, technique):
        q.put(('progress', msg, technique))

    def run_extract():
        try:
            text = extract_text_from_file(file_like, fname, model_id=model_id, on_progress=on_progress)
            q.put(('done', text, fname))
        except Exception as e:
            q.put(('error', str(e), ''))

    t = threading.Thread(target=run_extract)
    t.start()

    def generate():
        while True:
            try:
                item = q.get(timeout=120)
            except Exception:
                yield 'data: ' + json.dumps({'error': 'Timeout'}) + '\n\n'
                break
            if item[0] == 'progress':
                yield 'data: ' + json.dumps({'progress': item[1], 'technique': item[2]}) + '\n\n'
            elif item[0] == 'done':
                yield 'data: ' + json.dumps({'done': True, 'text': item[1], 'filename': item[2]}) + '\n\n'
                break
            elif item[0] == 'error':
                yield 'data: ' + json.dumps({'error': item[1]}) + '\n\n'
                break

    return Response(generate(), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/build', methods=['POST'])
def api_build():
    data = request.get_json(force=True)
    text = data.get('text', '').strip()
    if not text:
        return json_response({'error': 'No text provided'}, 400)

    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        'status': 'starting',
        'current_chunk': 0,
        'total_chunks': 0,
        'graph': None,
        'graph_data': {'nodes': [], 'edges': [], 'typeColors': {}},
        'source_text': text,
        'messages': [],
        'errors': []
    }

    model_id = data.get('model', None)
    domain = data.get('domain', 'auto')
    thread = threading.Thread(target=build_graph_async, args=(session_id, text, model_id, domain))
    thread.daemon = True
    thread.start()

    return json_response({'session_id': session_id})


@app.route('/api/status/<session_id>')
def api_status(session_id):
    session = sessions.get(session_id)
    if not session:
        return json_response({'error': 'Session not found'}, 404)
    return json_response({
        'status': session['status'],
        'current_chunk': session['current_chunk'],
        'total_chunks': session['total_chunks'],
        'graph_data': session['graph_data'],
        'errors': session.get('errors', []),
        'schema': session.get('schema', None),
        'summary': session.get('summary', None),
        'report': session.get('report', None)
    })


@app.route('/api/stop/<session_id>', methods=['POST'])
def api_stop(session_id):
    session = sessions.get(session_id)
    if not session:
        return json_response({'error': 'Session not found'}, 404)
    session['stopped'] = True
    return json_response({'ok': True})


@app.route('/api/report/<session_id>', methods=['POST'])
def api_report(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    if session.get('report'):
        return json_response({'report': session['report']})
    model_id = session.get('model_id')
    report = generate_report(session, model_id)
    session['report'] = report
    return json_response({'report': report})


MAX_HISTORY_TURNS = 10
MAX_CONTEXT_CHARS = 6000


def build_qa_system_prompt(session):
    ctx = get_full_context(session['graph'])
    source = session['source_text']
    if len(source) > MAX_CONTEXT_CHARS:
        source = source[:MAX_CONTEXT_CHARS] + '\n[... truncated ...]'
    return """You are a research assistant for the Orion platform. Answer questions using ONLY the provided data.
If the answer is not in the data, say so. Be concise and professional.
When referencing entities or relationships, be specific about names and types.
If a follow-up question references something from the conversation, use context from previous exchanges.

Entity & Relationship Data:
""" + ctx + """

Source Text:
""" + source


@app.route('/api/ask/<session_id>', methods=['POST'])
def api_ask(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    data = request.get_json(force=True)
    question = data.get('question', '').strip()
    if not question:
        return json_response({'error': 'No question provided'}, 400)

    model_id = session.get('model_id')
    system_prompt = build_qa_system_prompt(session)
    history = session.get('messages', [])
    window = history[-(MAX_HISTORY_TURNS * 2):]

    msgs = []
    for msg in window:
        msgs.append({'role': msg['role'], 'content': msg['content']})

    answer = llm_complete(question, model_id=model_id, system_prompt=system_prompt, messages=msgs)

    session['messages'].append({'role': 'user', 'content': question})
    session['messages'].append({'role': 'assistant', 'content': answer})

    return json_response({'answer': answer})


@app.route('/api/ask_stream/<session_id>', methods=['POST'])
def api_ask_stream(session_id):
    import time
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    data = request.get_json(force=True)
    question = data.get('question', '').strip()
    if not question:
        return json_response({'error': 'No question provided'}, 400)

    model_id = session.get('model_id')
    system_prompt = build_qa_system_prompt(session)
    history = session.get('messages', [])
    window = history[-(MAX_HISTORY_TURNS * 2):]
    msgs = [{'role': m['role'], 'content': m['content']} for m in window]

    answer = llm_complete(question, model_id=model_id, system_prompt=system_prompt, messages=msgs)

    session['messages'].append({'role': 'user', 'content': question})
    session['messages'].append({'role': 'assistant', 'content': answer})

    def generate():
        words = answer.split(' ')
        for i, word in enumerate(words):
            chunk = word + (' ' if i < len(words) - 1 else '')
            yield 'data: ' + json.dumps({'chunk': chunk}) + '\n\n'
            time.sleep(0.03)
        yield 'data: ' + json.dumps({'done': True}) + '\n\n'

    return Response(generate(), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/whatif/<session_id>', methods=['POST'])
def api_whatif(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    data = request.get_json(force=True)
    scenario = data.get('scenario', '').strip()
    if not scenario:
        return json_response({'error': 'No scenario provided'}, 400)

    model_id = session.get('model_id')
    ctx = get_full_context(session['graph'])
    source = session['source_text']
    if len(source) > MAX_CONTEXT_CHARS:
        source = source[:MAX_CONTEXT_CHARS] + '\n[... truncated ...]'

    prompt = WHATIF_SIMULATE_PROMPT.format(context=ctx, source=source, scenario=scenario)
    resp = llm_complete(prompt, model_id=model_id)
    raw = clean_llm_json(resp)

    try:
        result = json.loads(raw)
    except Exception:
        result = {
            'summary': resp, 'overall_risk': 'medium', 'risk_score': 50,
            'simulation_rounds': [], 'agent_perspectives': [],
            'broken_relationships': [], 'new_relationships': [], 'recommendations': []
        }

    if 'whatif_history' not in session:
        session['whatif_history'] = []
    session['whatif_history'].append({'scenario': scenario, 'result': result})

    return json_response({'result': result})


@app.route('/api/whatif_interview/<session_id>', methods=['POST'])
def api_whatif_interview(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    data = request.get_json(force=True)
    entity_name = data.get('entity', '').strip()
    question = data.get('question', '').strip()
    scenario = data.get('scenario', '').strip()
    agent_reaction = data.get('reaction', '')

    if not entity_name or not question:
        return json_response({'error': 'Entity and question required'}, 400)

    G = session['graph']
    model_id = session.get('model_id')

    entity_context_lines = []
    node_data = G.nodes.get(entity_name, {})
    entity_type = node_data.get('type', 'Entity')
    entity_context_lines.append('Type: ' + entity_type)
    if node_data.get('description'):
        entity_context_lines.append('Description: ' + node_data['description'])
    entity_context_lines.append('')
    entity_context_lines.append('Relationships:')
    for src, tgt, edata in G.edges(data=True):
        if src == entity_name:
            entity_context_lines.append('  You --[' + edata.get('relation', '') + ']--> ' + tgt)
        elif tgt == entity_name:
            entity_context_lines.append('  ' + src + ' --[' + edata.get('relation', '') + ']--> You')

    prompt = WHATIF_INTERVIEW_PROMPT.format(
        entity_name=entity_name, entity_type=entity_type,
        entity_context='\n'.join(entity_context_lines),
        scenario=scenario, agent_reaction=agent_reaction, question=question
    )
    answer = llm_complete(prompt, model_id=model_id)
    return json_response({'answer': answer.strip(), 'entity': entity_name})


@app.route('/api/whatif_suggestions/<session_id>')
def api_whatif_suggestions(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)

    G = session['graph']
    suggestions = []
    degrees = sorted(G.degree(), key=lambda x: x[1], reverse=True)

    if len(degrees) >= 1:
        top = degrees[0][0]
        suggestions.append('What if ' + top + ' is removed from the system?')
    if len(degrees) >= 2:
        second = degrees[1][0]
        suggestions.append('What if ' + second + ' changes its role entirely?')

    rel_counts = {}
    for _, _, data in G.edges(data=True):
        r = data.get('relation', '')
        rel_counts[r] = rel_counts.get(r, 0) + 1
    if rel_counts:
        top_rel = max(rel_counts, key=rel_counts.get)
        suggestions.append('What if all ' + top_rel + ' relationships are dissolved?')

    if len(degrees) >= 3:
        a = degrees[0][0]
        b = degrees[2][0]
        suggestions.append('What if ' + a + ' and ' + b + ' merge into one?')
    if len(degrees) >= 4:
        c = degrees[1][0]
        d = degrees[3][0]
        if not G.has_edge(c, d) and not G.has_edge(d, c):
            suggestions.append('What if ' + c + ' starts reporting to ' + d + '?')

    return json_response({'suggestions': suggestions})


@app.route('/api/connect/<session_id>', methods=['POST'])
def api_connect(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    data = request.get_json(force=True)
    source = data.get('source', '').strip()
    target = data.get('target', '').strip()
    if not source or not target:
        return json_response({'error': 'Source and target entities required'}, 400)

    G = session['graph']
    if source not in G.nodes:
        return json_response({'error': 'Entity not found: ' + source}, 404)
    if target not in G.nodes:
        return json_response({'error': 'Entity not found: ' + target}, 404)

    U = G.to_undirected()
    try:
        path_nodes = nx.shortest_path(U, source, target)
    except nx.NetworkXNoPath:
        return json_response({'error': 'No connection found between ' + source + ' and ' + target, 'no_path': True}, 200)

    path_edges = []
    for i in range(len(path_nodes) - 1):
        a, b = path_nodes[i], path_nodes[i + 1]
        if G.has_edge(a, b):
            edata = G.edges[a, b]
            path_edges.append({'from': a, 'to': b, 'relation': edata.get('relation', 'RELATED_TO'), 'description': edata.get('description', ''), 'direction': 'forward'})
        elif G.has_edge(b, a):
            edata = G.edges[b, a]
            path_edges.append({'from': b, 'to': a, 'relation': edata.get('relation', 'RELATED_TO'), 'description': edata.get('description', ''), 'direction': 'reverse'})
        else:
            path_edges.append({'from': a, 'to': b, 'relation': 'CONNECTED', 'description': '', 'direction': 'undirected'})

    path_desc_lines = []
    for i, edge in enumerate(path_edges):
        path_desc_lines.append(str(i + 1) + '. ' + edge['from'] + ' --[' + edge['relation'] + ']--> ' + edge['to'] + (' (' + edge['description'] + ')' if edge['description'] else ''))
    for node_name in path_nodes:
        ndata = G.nodes[node_name]
        path_desc_lines.append('  Entity: ' + node_name + ' (type: ' + ndata.get('type', 'Unknown') + ') — ' + ndata.get('description', ''))

    model_id = session.get('model_id')
    ctx = get_full_context(G)
    if len(ctx) > MAX_CONTEXT_CHARS:
        ctx = ctx[:MAX_CONTEXT_CHARS] + '\n[... truncated ...]'

    prompt = CONNECT_PROMPT.format(path_description='\n'.join(path_desc_lines), context=ctx)
    resp = llm_complete(prompt, model_id=model_id)
    raw = clean_llm_json(resp)

    try:
        analysis = json.loads(raw)
    except Exception:
        analysis = {'hops': [], 'synthesis': resp, 'implications': []}

    return json_response({'path': path_nodes, 'edges': path_edges, 'analysis': analysis, 'hop_count': len(path_edges)})


@app.route('/api/entities/<session_id>')
def api_entities_list(session_id):
    session = sessions.get(session_id)
    if not session or not session.get('graph'):
        return json_response({'error': 'No graph built yet'}, 400)
    G = session['graph']
    entities = [{'name': node, 'type': data.get('type', 'Unknown')} for node, data in G.nodes(data=True)]
    entities.sort(key=lambda x: x['name'])
    return json_response({'entities': entities})


if __name__ == '__main__':
    _init_models()
    print('\n  Orion — Local Development Server')
    print('  ================================')
    if AZURE_ENDPOINT:
        print('  LLM: Azure OpenAI (' + AZURE_ENDPOINT + ')')
    elif OPENAI_API_KEY:
        print('  LLM: OpenAI')
    else:
        print('  WARNING: No API key set! Export OPENAI_API_KEY or AZURE_OPENAI_API_KEY + AZURE_OPENAI_ENDPOINT')
    print('  Default model: ' + DEFAULT_MODEL)
    print('  Models available: ' + ', '.join(AVAILABLE_MODELS))
    print('  Server: http://localhost:5001')
    print()
    app.run(host='0.0.0.0', port=5001, debug=True, threaded=True)
